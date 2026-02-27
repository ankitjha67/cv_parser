import os
import re
from pathlib import Path
from typing import Optional
import PyPDF2
from docx import Document

def load_document(file_path: str, file_content: Optional[bytes] = None) -> str:
    """
    Load document from file path or bytes content.
    Supports TXT, PDF, DOCX formats.
    """
    if file_content:
        # Determine file type from extension
        ext = Path(file_path).suffix.lower()
        if ext == '.pdf':
            return _parse_pdf_bytes(file_content)
        elif ext in ['.docx', '.doc']:
            return _parse_docx_bytes(file_content)
        else:
            return file_content.decode('utf-8', errors='ignore')
    else:
        ext = Path(file_path).suffix.lower()
        if ext == '.pdf':
            return _parse_pdf_file(file_path)
        elif ext in ['.docx', '.doc']:
            return _parse_docx_file(file_path)
        else:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()

def _parse_pdf_file(path: str) -> str:
    """Parse PDF file to text."""
    try:
        with open(path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            text = ''
            for page in reader.pages:
                text += page.extract_text() + '\n'
            return clean_text(text)
    except Exception as e:
        raise ValueError(f"Failed to parse PDF: {str(e)}")

def _parse_pdf_bytes(content: bytes) -> str:
    """Parse PDF from bytes to text."""
    try:
        import io
        reader = PyPDF2.PdfReader(io.BytesIO(content))
        text = ''
        for page in reader.pages:
            text += page.extract_text() + '\n'
        return clean_text(text)
    except Exception as e:
        raise ValueError(f"Failed to parse PDF: {str(e)}")

def _parse_docx_file(path: str) -> str:
    """Parse DOCX file to text."""
    try:
        doc = Document(path)
        text = '\n'.join([para.text for para in doc.paragraphs])
        return clean_text(text)
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {str(e)}")

def _parse_docx_bytes(content: bytes) -> str:
    """Parse DOCX from bytes to text."""
    try:
        import io
        doc = Document(io.BytesIO(content))
        text = '\n'.join([para.text for para in doc.paragraphs])
        return clean_text(text)
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {str(e)}")

def clean_text(text: str) -> str:
    """Clean and normalize text."""
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)
    # Remove special characters but keep bullets
    text = re.sub(r'[^\w\s\-•●○◦▪▫⁃‣⦿⦾.,:;()\/\n@]', '', text)
    # Normalize line breaks
    text = re.sub(r'\n\s*\n', '\n\n', text)
    return text.strip()
